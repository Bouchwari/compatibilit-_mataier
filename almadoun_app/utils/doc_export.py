# -*- coding: utf-8 -*-
"""
utils/doc_export.py
Fills Word templates with data from the database and saves them.
Templates are in the  templet/  folder next to this package.
"""

import os
from datetime import datetime
from docx import Document

BASE_DIR   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLET_DIR = os.path.join(BASE_DIR, "templet")
EXPORTS_DIR = os.path.join(BASE_DIR, "exports")


def _ensure_exports():
    os.makedirs(EXPORTS_DIR, exist_ok=True)


def _set_cell_text(cell, text: str):
    """Overwrite cell text, keeping the first paragraph's first run."""
    for para in cell.paragraphs:
        # Clear all runs
        for run in para.runs:
            run.text = ""
        # Write into first run; if no runs exist, just set paragraph text
        if para.runs:
            para.runs[0].text = str(text)
        else:
            para.clear()
            para.add_run(str(text))
        break   # Only touch first paragraph


def _replace_in_doc(doc: Document, mapping: dict):
    """Replace placeholder keys with values everywhere in the document,
    including text boxes (w:txbxContent inside drawings/shapes)."""

    def _fix_para(para):
        for key, val in mapping.items():
            if key in para.text:
                # First pass: try direct run replacement
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(val))
                # Second pass (fallback): key might be split across runs
                if key in para.text:
                    full = para.text
                    for k, v in mapping.items():
                        full = full.replace(k, str(v))
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = full

    # Regular paragraphs
    for para in doc.paragraphs:
        _fix_para(para)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _fix_para(para)

    # Text boxes (w:txbxContent) — these contain "شهادة تسلم رقم" and "حرر بـ"
    from docx.oxml.ns import qn
    TXBX = qn("w:txbxContent")
    W_P  = qn("w:p")
    W_R  = qn("w:r")
    W_T  = qn("w:t")

    for txbx in doc.element.body.iter(TXBX):
        for wp in txbx.iter(W_P):
            # Build full paragraph text from runs
            runs = list(wp.iter(W_R))
            full_text = "".join(
                (wt.text or "") for r in runs for wt in r.iter(W_T)
            )
            needs_rewrite = any(k in full_text for k in mapping)
            if needs_rewrite:
                new_text = full_text
                for k, v in mapping.items():
                    new_text = new_text.replace(k, str(v))
                # Write result into first <w:t>, clear the rest
                all_wt = [wt for r in runs for wt in r.iter(W_T)]
                for i, wt in enumerate(all_wt):
                    wt.text = new_text if i == 0 else ""


# ═══════════════════════════════════════════════════════════════════
#  شهادة التسلم  →  Word
# ═══════════════════════════════════════════════════════════════════
def export_certificate_docx(cert_num, recipient, position, date, items):
    """
    items: list of dicts with keys 'item', 'unit', 'qty'
    Returns the output file path.
    """
    _ensure_exports()
    template_path = os.path.join(TEMPLET_DIR, "شهادة تسلم -.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # ── Load school settings ────────────────────────────────────────
    try:
        import sys as _sys
        _sys.path.insert(0, BASE_DIR)
        import database as _db
        _s          = _db.get_settings()
        school_name = _s.get("school_name", "") or ""
        academy     = _s.get("academy",     "") or ""
        delegation  = _s.get("delegation",  "") or ""
        address     = _s.get("address",     "") or ""
        logo_path   = _s.get("logo_path",   "") or ""
    except Exception:
        school_name = academy = delegation = address = logo_path = ""

    # ── Format date as DD/MM/YYYY ───────────────────────────────────
    try:
        from datetime import datetime as _dt
        _d      = _dt.strptime(str(date), "%Y-%m-%d")
        date_ar = _d.strftime("%d/%m/%Y")
    except Exception:
        date_ar = str(date)

    # ── Build mapping using ACTUAL template placeholder text ────────
    # (found by scanning the template XML — see docs/template_scan.txt)
    mapping = {
        # Certificate number (in text box: "شهادة تسلم " + "رقم ......")
        "رقم ......":                   f"رقم {cert_num}",

        # Beneficiary name (Arabic ellipsis dots after السيد(ة):)
        "………………………….":              str(recipient),

        # Position / role
        "............................................": str(position),

        # حرر بـ city (18 dots after حرر بـ)
        "..................":             str(address) if address else "..................",

        # The full حرر بـ date line (fallback for fragmented runs)
        ".../.../…….":                  date_ar,

        # School header lines (replace hard-coded school name)
        "الثانوية الإعدادية ألمدون":   school_name if school_name else "الثانوية الإعدادية ألمدون",
    }

    # Only replace academy/delegation if user filled them in settings
    if academy:
        mapping["درعة تافيلالت"] = academy
    if delegation:
        mapping["تنغير"] = delegation

    _replace_in_doc(doc, mapping)

    # ── Insert logo at the very top ─────────────────────────────────
    if logo_path and os.path.isfile(logo_path):
        try:
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            logo_para = doc.paragraphs[0].insert_paragraph_before()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_para.add_run().add_picture(logo_path, width=Inches(1.2))
        except Exception:
            pass  # never crash the export over a logo

    # ── Fill items table ─────────────────────────────────────────────
    items_table = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "تعيين" in cell.text or "مادة" in cell.text:
                    items_table = table
                    break
            if items_table:
                break
        if items_table:
            break

    if items_table:
        data_rows = []
        for row in items_table.rows:
            texts = [c.text.strip() for c in row.cells]
            if all(t == "" or t.isdigit() for t in texts):
                data_rows.append(row)

        for i, item in enumerate(items[:len(data_rows)]):
            cells = data_rows[i].cells
            try:
                if len(cells) >= 4:
                    _set_cell_text(cells[0], str(i + 1))
                    _set_cell_text(cells[1], str(item.get("item", "")))
                    _set_cell_text(cells[2], str(item.get("unit", "")))
                    _set_cell_text(cells[3], str(item.get("qty", "")))
                elif len(cells) >= 3:
                    _set_cell_text(cells[0], str(i + 1))
                    _set_cell_text(cells[1], str(item.get("item", "")))
                    _set_cell_text(cells[2], str(item.get("qty", "")))
            except Exception:
                pass

    date_safe = str(date).replace("-", "")
    out_path = os.path.join(EXPORTS_DIR, f"شهادة_التسلم_{cert_num}_{date_safe}.docx")
    doc.save(out_path)
    return out_path


# ═══════════════════════════════════════════════════════════════════
#  بطاقة المخزون (فيش المادة)  →  Word
# ═══════════════════════════════════════════════════════════════════
def export_stock_card_docx(item_name, unit, movements):
    """
    movements: list of dicts with keys:
        date, type ('دخول'/'خروج'), quantity, beneficiary/who, balance, notes

    Template column layout (8 cols, rows 0-2 are headers):
      col[0]  → Imputation / النفقة        — leave blank
      col[1]  → Date de livraison / التاريخ — date
      col[2]  → Entrée / دخل               — in quantity
      col[3]  → Sortie / خرج               — out quantity
      col[4]  → Reste / الباقي             — running balance
      col[5]  → Prénom & Nom               — beneficiary name
      col[6]  → Profession / المهنة        — leave blank
      col[7]  → (extra merged cell)

    Item name: row 0 cell[0] run — text "Matière: ..."
    """
    _ensure_exports()
    template_path = os.path.join(TEMPLET_DIR, "fiche_stock.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    table = doc.tables[0]   # Only one table in this template

    # ── 1. Set item name in the header cell (row 0, cell 0) ─────────
    header_cell = table.rows[0].cells[0]
    name_set = False
    for para in header_cell.paragraphs:
        for run in para.runs:
            if run.text.strip():
                run.text = f"Matière: {item_name}"
                # Clear remaining runs in this paragraph
                para_runs = para.runs
                for extra_run in para_runs[1:]:
                    extra_run.text = ""
                name_set = True
                break
        if name_set:
            break

    # ── 2. Data rows: template rows 3..34 (32 available lines) ──────
    # Row 3 contains sample data — overwrite ALL rows from index 3 onward
    data_rows = list(table.rows[3:])

    for i, row in enumerate(data_rows):
        cells = row.cells
        num_cells = len(cells)
        if i < len(movements):
            m   = movements[i]
            qty = float(m.get("quantity", 0))
            bal = float(m.get("balance", 0))

            in_v  = str(int(qty))  if m.get("type") == "دخول" else ""
            out_v = str(int(qty))  if m.get("type") == "خروج" else ""
            who   = str(m.get("beneficiary") or m.get("who") or "")
            date  = str(m.get("date", ""))

            if num_cells > 0: _set_cell_text(cells[0], "")           # Imputation
            if num_cells > 1: _set_cell_text(cells[1], date)          # Date
            if num_cells > 2: _set_cell_text(cells[2], in_v)          # Entrée
            if num_cells > 3: _set_cell_text(cells[3], out_v)         # Sortie
            if num_cells > 4: _set_cell_text(cells[4], str(int(bal))) # Reste
            if num_cells > 5: _set_cell_text(cells[5], who)           # Nom
            if num_cells > 6: _set_cell_text(cells[6], "")            # Profession
        else:
            # Clear any pre-filled example data beyond our movements
            for ci in range(min(8, num_cells)):
                _set_cell_text(cells[ci], "")

    name_safe = item_name.replace(" ", "_")[:30]
    date_safe = datetime.now().strftime("%Y%m%d_%H%M")
    out_path  = os.path.join(EXPORTS_DIR, f"فيش_المادة_{name_safe}_{date_safe}.docx")
    doc.save(out_path)
    return out_path


# ═══════════════════════════════════════════════════════════════════
#  كشف المخزون العام  →  Excel
# ═══════════════════════════════════════════════════════════════════
def export_inventory_excel(dashboard_data):
    """
    dashboard_data: list of rows from db.get_dashboard_data()
    Returns the output file path.
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    _ensure_exports()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "كشف المخزون"
    ws.sheet_view.rightToLeft = True

    school = "الثانوية الإعدادية ألمدون — جماعة اغيل نومكون"
    ws.merge_cells("A1:H1")
    ws["A1"] = school
    ws["A1"].font = Font(bold=True, size=13, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor="0D2137")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 24

    ws.merge_cells("A2:H2")
    ws["A2"] = f"كشف المخزون العام — تاريخ الإصدار: {datetime.now().strftime('%Y-%m-%d')}"
    ws["A2"].font = Font(bold=True, size=11, color="FFFFFF")
    ws["A2"].fill = PatternFill("solid", fgColor="1A3A52")
    ws["A2"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 20

    headers = ["الفئة", "التسمية", "الوحدة", "الدخول", "الخروج", "الرصيد", "الحالة", "آخر عملية"]
    widths  = [18,      32,        12,        12,        12,        12,       16,        16]
    hdr_fill = PatternFill("solid", fgColor="2E5F8A")
    thin   = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, (h, w) in enumerate(zip(headers, widths), 1):
        cell = ws.cell(row=3, column=col, value=h)
        cell.font = Font(bold=True, size=10, color="FFFFFF")
        cell.fill = hdr_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border
        ws.column_dimensions[cell.column_letter].width = w
    ws.row_dimensions[3].height = 18

    fill_even = PatternFill("solid", fgColor="EAF2FA")
    fill_odd  = PatternFill("solid", fgColor="FFFFFF")
    fill_low  = PatternFill("solid", fgColor="FFF3CD")
    fill_out  = PatternFill("solid", fgColor="FFE0E0")

    for i, r in enumerate(dashboard_data):
        row_num = i + 4
        bal = r["balance"]

        if bal <= 0:
            status = "🔴 نفد المخزون"; rfill = fill_out
        elif bal <= 5:
            status = "🟡 منخفض";       rfill = fill_low
        else:
            status = "🟢 متوفر";        rfill = fill_even if i % 2 == 0 else fill_odd

        values = [
            r["category"], r["name"], r["unit"] or "",
            int(r["total_in"]), int(r["total_out"]), int(bal),
            status, r["last_date"] or "—"
        ]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row_num, column=col, value=val)
            cell.fill = rfill
            cell.border = border
            cell.alignment = Alignment(
                horizontal="right" if col <= 3 else "center",
                vertical="center")
            if col == 6:
                cell.font = Font(
                    bold=True,
                    color="008000" if bal > 5 else ("FF8C00" if bal > 0 else "CC0000"))
        ws.row_dimensions[row_num].height = 16

    ws.freeze_panes = "A4"
    date_safe = datetime.now().strftime("%Y%m%d_%H%M")
    out_path  = os.path.join(EXPORTS_DIR, f"كشف_المخزون_{date_safe}.xlsx")
    wb.save(out_path)
    return out_path


# ═══════════════════════════════════════════════════════════════════
#  قالب الجرد  →  Excel  (تصدير + استيراد)
# ═══════════════════════════════════════════════════════════════════
def export_inventory_template_excel(all_items):
    """
    Generate a fillable Excel inventory template pre-loaded with all items.
    all_items: list of dicts with keys name, unit, category (from db.get_dashboard_data())
    Returns the saved file path.
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.worksheet.datavalidation import DataValidation

    _ensure_exports()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "الجرد"
    ws.sheet_view.rightToLeft = True

    # ── TITLE ────────────────────────────────────────────────────────
    ws.merge_cells("A1:H1")
    ws["A1"] = "الثانوية الإعدادية ألمدون — قالب الجرد السنوي"
    ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor="0D2137")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    ws.merge_cells("A2:H2")
    ws["A2"] = ("⚠️  أدخل الكمية المستلمة في العمود E فقط. "
                "لا تغير أعمدة A - D. احفظ الملف ثم استورده في التطبيق.")
    ws["A2"].font = Font(bold=True, size=10, color="7B3F00")
    ws["A2"].fill = PatternFill("solid", fgColor="FFF3CD")
    ws["A2"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[2].height = 32

    # ── COLUMN HEADERS ────────────────────────────────────────────────
    headers = ["الفئة", "التسمية (اسم المادة)", "الوحدة",
               "الرصيد الحالي", "★ الكمية المستلمة ★",
               "تاريخ الاستلام", "المورد", "ملاحظات"]
    widths  = [20, 38, 14, 16, 22, 18, 22, 24]

    thin   = Side(style="thin", color="AAAAAA")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    h_fill = PatternFill("solid", fgColor="1A3A52")

    for col, (h, w) in enumerate(zip(headers, widths), 1):
        cell = ws.cell(row=3, column=col, value=h)
        cell.font = Font(bold=True, size=10, color="FFFFFF")
        cell.fill = h_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border
        ws.column_dimensions[cell.column_letter].width = w
    ws.row_dimensions[3].height = 20

    # ── DATA ROWS ─────────────────────────────────────────────────────
    cat_fills = {
        "مكتبيات":         PatternFill("solid", fgColor="EAF2FA"),
        "خراطيش الطباعة": PatternFill("solid", fgColor="FFF4E5"),
        "ادوات التدريس":  PatternFill("solid", fgColor="E8F5E9"),
        "ادوات النظافة":  PatternFill("solid", fgColor="F3E5F5"),
        "الصيانة":         PatternFill("solid", fgColor="FFF8E1"),
        "مختلفات":         PatternFill("solid", fgColor="F5F5F5"),
    }
    star_fill = PatternFill("solid", fgColor="FFFFCC")

    today = datetime.now().strftime("%Y-%m-%d")
    for i, item in enumerate(all_items):
        row_num = i + 4
        cat = item["category"] if "category" in item.keys() else str(item[2] if len(item) > 2 else "")
        row_fill = cat_fills.get(cat, PatternFill("solid", fgColor="FFFFFF"))

        try:
            name    = item["name"]
            unit    = item["unit"] or ""
            balance = int(item["balance"] or 0)
        except (IndexError, KeyError):
            # Fallback for positional sqlite3.Row: name=0, unit=1, category=2, total_in=3, total_out=4, balance=5
            name    = str(item[0])
            unit    = str(item[1] or "")
            balance = int(item[5] or 0)

        values = [
            cat,
            name,
            unit,
            balance,
            "",          # ← user fills this
            today,       # ← default date (user can change)
            "",          # ← supplier
            "",          # ← notes
        ]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row_num, column=col, value=val)
            cell.border = border
            cell.alignment = Alignment(horizontal="right", vertical="center")
            if col == 5:  # quantity column — highlighted
                cell.fill = star_fill
                cell.font = Font(bold=True, size=11)
            elif col == 4:  # current balance — non-editable visual
                cell.font = Font(italic=True, color="888888")
                cell.fill = PatternFill("solid", fgColor="F0F0F0")
            else:
                cell.fill = row_fill
        ws.row_dimensions[row_num].height = 16

    # Lock cols A-D visually (gray) and protect sheet
    ws.protection.sheet = False   # keep editable; just styling guides the user

    ws.freeze_panes = "E4"

    date_safe = datetime.now().strftime("%Y%m%d_%H%M")
    out_path = os.path.join(EXPORTS_DIR, f"قالب_الجرد_{date_safe}.xlsx")
    wb.save(out_path)
    return out_path


def import_inventory_from_excel(file_path):
    """
    Read a filled inventory template Excel file and insert movements into DB.
    Returns (imported_count, skipped_list, errors_list).
    Columns (1-indexed): A=category, B=name, C=unit, D=balance(ignore),
                         E=qty_received, F=date, G=supplier, H=notes
    """
    import openpyxl
    import sys, os
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    import database as db

    wb = openpyxl.load_workbook(file_path, data_only=True)
    ws = wb.active

    imported = 0
    skipped  = []
    errors   = []

    for row in ws.iter_rows(min_row=4, values_only=True):
        cat   = str(row[0] or "").strip()
        name  = str(row[1] or "").strip()
        unit  = str(row[2] or "").strip()
        # col 3 = current balance — ignored
        qty_raw = row[4]  # column E
        date_raw = row[5]  # column F
        supplier = str(row[6] or "").strip()
        notes    = str(row[7] or "").strip()

        if not name:
            continue   # empty row

        try:
            qty = float(qty_raw or 0)
        except (TypeError, ValueError):
            qty = 0

        if qty <= 0:
            skipped.append(name)
            continue

        # Parse date
        if hasattr(date_raw, "strftime"):
            date = date_raw.strftime("%Y-%m-%d")
        elif date_raw:
            date = str(date_raw)[:10]
        else:
            date = datetime.now().strftime("%Y-%m-%d")

        try:
            # Create item if not in DB
            existing = [r["name"] for r in db.get_items_by_category(cat)] if cat else []
            if name not in existing:
                db.add_item(name, unit or "قطعة", cat or "مختلفات")

            # Add دخول movement
            db.add_movement(
                item_name=name,
                category=cat or "مختلفات",
                date=date,
                mov_type="دخول",
                quantity=qty,
                beneficiary=supplier,
                notes=notes or "استيراد من الجرد"
            )
            imported += 1
        except Exception as e:
            errors.append(f"{name}: {e}")

    return imported, skipped, errors


# ═══════════════════════════════════════════════════════════════════
#  تصدير جميع بطاقات المخزون → ملف Word واحد جاهز للطباعة
# ═══════════════════════════════════════════════════════════════════
def _page_break_element():
    """Return an XML paragraph element containing a hard page break."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p  = OxmlElement("w:p")
    r  = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def export_all_stock_cards(progress_callback=None):
    """
    Export ALL item stock cards into ONE Word document, one fiche per page.
    Returns (output_path, count_exported, skipped_items).
    """
    from copy import deepcopy
    import sys
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    import database as db

    _ensure_exports()

    all_items = db.get_dashboard_data()
    total     = len(all_items)
    exported  = 0
    skipped   = []
    tmp_files = []   # individual temp files to merge then delete

    # ── Step 1: generate individual temp files ────────────────────
    for idx, item in enumerate(all_items):
        name = item["name"]
        unit = item["unit"] or ""

        if progress_callback:
            progress_callback(idx + 1, total, name)

        movements = db.get_movements(item_name=name)
        if not movements:
            skipped.append(name)
            continue

        # Build rows with running balance (chronological order)
        rows = []
        bal  = 0
        for m in reversed(movements):
            qty = float(m["quantity"])
            bal += qty if m["type"] == "دخول" else -qty
            rows.append({
                "date":     m["date"] or "",
                "who":      m["beneficiary"] or "",
                "type":     m["type"],
                "quantity": qty,
                "balance":  bal,
                "notes":    m["notes"] or "",
            })
        rows.reverse()

        try:
            tmp_path = export_stock_card_docx(name, unit, rows)
            tmp_files.append(tmp_path)
            exported += 1
        except Exception as e:
            skipped.append(f"{name} ({e})")

    if not tmp_files:
        # Nothing to merge
        return None, 0, skipped

    # ── Step 2: merge into ONE document ───────────────────────────
    master = Document(tmp_files[0])

    for path in tmp_files[1:]:
        # Insert hard page break before each new fiche
        master.element.body.append(_page_break_element())

        sub = Document(path)
        body_elements = list(sub.element.body)
        for el in body_elements:
            # Skip final section properties (sectPr) — keep master's layout
            tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
            if tag == "sectPr":
                continue
            master.element.body.append(deepcopy(el))

    # ── Step 3: save merged file ──────────────────────────────────
    date_safe  = datetime.now().strftime("%Y%m%d_%H%M")
    out_path   = os.path.join(EXPORTS_DIR, f"جميع_بطاقات_المخزون_{date_safe}.docx")
    master.save(out_path)

    # ── Step 4: clean up temp files ───────────────────────────────
    for f in tmp_files:
        try:
            os.remove(f)
        except Exception:
            pass

    return out_path, exported, skipped
